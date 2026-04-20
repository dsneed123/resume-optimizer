from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


import re as _re

_INLINE_MD_RE = _re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')


def _add_inline_md_runs(para, text: str, font_name: str, font_size: float):
    """Split text on **bold** and *italic* markers, adding styled runs."""
    pos = 0
    for m in _INLINE_MD_RE.finditer(text):
        if m.start() > pos:
            _add_run(para, text[pos:m.start()], font_name, font_size)
        if m.group(2) is not None:
            _add_run(para, m.group(2), font_name, font_size, bold=True)
        else:
            _add_run(para, m.group(3), font_name, font_size, italic=True)
        pos = m.end()
    if pos < len(text):
        _add_run(para, text[pos:], font_name, font_size)


_FONT_MAP = {
    "Helvetica": "Arial",
    "Times New Roman": "Times New Roman",
    "Georgia": "Georgia",
    "Courier": "Courier New",
    "Verdana": "Verdana",
    "Calibri": "Calibri",
}


def _resolve_font(font_family: str) -> str:
    return _FONT_MAP.get(font_family, font_family)


def _set_paragraph_spacing(para, space_before_pt: float, space_after_pt: float, line_height: float):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing = line_height


def _add_horizontal_rule(para, style: str = "thin"):
    """Insert a bottom border on a paragraph to simulate a section divider."""
    if style == "none":
        return
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    if style == "thick":
        bottom.set(qn("w:val"), "thick")
        bottom.set(qn("w:sz"), "18")
    elif style == "double":
        bottom.set(qn("w:val"), "double")
        bottom.set(qn("w:sz"), "6")
    else:
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(para, text: str, font_name: str, font_size: float, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _add_section_header(doc, title: str, font_name: str, font_size: float, section_spacing: float, divider_style: str = "thin"):
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, section_spacing, 2, 1.0)
    _add_run(para, title.upper(), font_name, font_size, bold=True)
    _add_horizontal_rule(para, divider_style)
    return para


def export_docx(resume_data: dict, typography: dict) -> bytes:
    """Render resume_data + typography to a single-page ATS-friendly DOCX."""
    t = typography
    font_name = _resolve_font(t.get("font_family", "Arial"))
    size_name = float(t.get("font_size_name", 20))
    size_header = float(t.get("font_size_section_header", 12))
    size_body = float(t.get("font_size_body", 10))
    size_detail = float(t.get("font_size_detail", 9))
    line_height = float(t.get("line_height", 1.15))
    para_spacing = float(t.get("paragraph_spacing", 4))
    section_spacing = float(t.get("section_spacing", 10))
    margin_top = float(t.get("margin_top", 0.5))
    margin_bottom = float(t.get("margin_bottom", 0.5))
    margin_left = float(t.get("margin_left", 0.6))
    margin_right = float(t.get("margin_right", 0.6))
    divider_style = t.get("section_divider_style", "thin")
    bullet_style = t.get("bullet_style", "filled")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(margin_top)
    section.bottom_margin = Inches(margin_bottom)
    section.left_margin = Inches(margin_left)
    section.right_margin = Inches(margin_right)

    # --- Header ---
    header_data = resume_data.get("header", {})
    name = header_data.get("name", "")
    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(name_para, 0, para_spacing, 1.0)
        _add_run(name_para, name, font_name, size_name, bold=True)

    contact_parts = [
        header_data.get("email", ""),
        header_data.get("phone", ""),
        header_data.get("location", ""),
        header_data.get("linkedin", ""),
        header_data.get("website", ""),
    ]
    contact_line = "  |  ".join(p for p in contact_parts if p)
    if contact_line:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(contact_para, 0, para_spacing, line_height)
        _add_run(contact_para, contact_line, font_name, size_detail)

    # --- Summary ---
    summary = resume_data.get("summary")
    if summary:
        _add_section_header(doc, "Summary", font_name, size_header, section_spacing, divider_style)
        para = doc.add_paragraph()
        _set_paragraph_spacing(para, 2, para_spacing, line_height)
        _add_run(para, summary, font_name, size_body)

    # --- Experience ---
    experience = resume_data.get("experience", [])
    if experience:
        _add_section_header(doc, "Experience", font_name, size_header, section_spacing, divider_style)
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            company = exp.get("company", "")
            title = exp.get("title", "")
            location = exp.get("location", "")
            start_date = exp.get("start_date", "")
            end_date = exp.get("end_date", "")

            if company or title:
                job_para = doc.add_paragraph()
                _set_paragraph_spacing(job_para, 4, 0, 1.0)
                if title:
                    _add_run(job_para, title, font_name, size_body, bold=True)
                if company:
                    sep = "  —  " if title else ""
                    _add_run(job_para, sep + company, font_name, size_body)

                date_loc_parts = [p for p in [location, _date_range(start_date, end_date)] if p]
                if date_loc_parts:
                    detail_para = doc.add_paragraph()
                    _set_paragraph_spacing(detail_para, 0, 2, 1.0)
                    _add_run(detail_para, "  |  ".join(date_loc_parts), font_name, size_detail, italic=True)

            _BULLET_CHARS = {"open": "\u25e6 ", "dash": "\u2013 ", "none": ""}
            for bullet in exp.get("bullets", []):
                if not bullet:
                    continue
                if bullet_style == "filled":
                    bul_para = doc.add_paragraph(style="List Bullet")
                else:
                    bul_para = doc.add_paragraph()
                    bul_para.paragraph_format.left_indent = Inches(0.25)
                    prefix = _BULLET_CHARS.get(bullet_style, "")
                    if prefix:
                        _add_run(bul_para, prefix, font_name, size_body)
                _set_paragraph_spacing(bul_para, 0, 1, line_height)
                _add_inline_md_runs(bul_para, str(bullet), font_name, size_body)

    # --- Education ---
    education = resume_data.get("education", [])
    if education:
        _add_section_header(doc, "Education", font_name, size_header, section_spacing, divider_style)
        for edu in education:
            if not isinstance(edu, dict):
                continue
            school = edu.get("school", "")
            degree = edu.get("degree", "")
            field_of_study = edu.get("field", "")
            graduation_date = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")
            honors = edu.get("honors", "")

            if school:
                sch_para = doc.add_paragraph()
                _set_paragraph_spacing(sch_para, 4, 0, 1.0)
                _add_run(sch_para, school, font_name, size_body, bold=True)

                degree_parts = [p for p in [degree, field_of_study] if p]
                detail_parts = [", ".join(degree_parts)]
                if graduation_date:
                    detail_parts.append(graduation_date)
                if gpa:
                    detail_parts.append(f"GPA: {gpa}")
                if honors:
                    detail_parts.append(honors)

                detail_line = "  |  ".join(p for p in detail_parts if p)
                if detail_line:
                    det_para = doc.add_paragraph()
                    _set_paragraph_spacing(det_para, 0, 2, 1.0)
                    _add_run(det_para, detail_line, font_name, size_detail)

    # --- Skills ---
    skills = resume_data.get("skills", [])
    if skills:
        _add_section_header(doc, "Skills", font_name, size_header, section_spacing, divider_style)
        for skill in skills:
            if not isinstance(skill, dict):
                continue
            category = skill.get("category", "")
            items = skill.get("items", [])
            if not items:
                continue
            skill_para = doc.add_paragraph()
            _set_paragraph_spacing(skill_para, 0, 2, line_height)
            if category:
                _add_run(skill_para, category + ": ", font_name, size_body, bold=True)
            _add_run(skill_para, ", ".join(str(i) for i in items), font_name, size_body)

    # --- Certifications ---
    certifications = resume_data.get("certifications", [])
    if certifications:
        _add_section_header(doc, "Certifications", font_name, size_header, section_spacing, divider_style)
        for cert in certifications:
            if not isinstance(cert, dict):
                continue
            cert_name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            if cert_name:
                cert_para = doc.add_paragraph()
                _set_paragraph_spacing(cert_para, 0, 2, line_height)
                _add_run(cert_para, cert_name, font_name, size_body, bold=True)
                detail_parts = [p for p in [issuer, date] if p]
                if detail_parts:
                    _add_run(cert_para, "  —  " + "  |  ".join(detail_parts), font_name, size_detail)

    # --- Projects ---
    projects = resume_data.get("projects", [])
    if projects:
        _add_section_header(doc, "Projects", font_name, size_header, section_spacing, divider_style)
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            proj_name = proj.get("name", "")
            description = proj.get("description", "")
            technologies = proj.get("technologies", "")
            url = proj.get("url", "")
            if proj_name:
                proj_para = doc.add_paragraph()
                _set_paragraph_spacing(proj_para, 4, 0, 1.0)
                _add_run(proj_para, proj_name, font_name, size_body, bold=True)
                if technologies:
                    _add_run(proj_para, f"  ({technologies})", font_name, size_detail)
                if description:
                    desc_para = doc.add_paragraph()
                    _set_paragraph_spacing(desc_para, 0, 2, line_height)
                    _add_run(desc_para, description, font_name, size_body)
                if url:
                    url_para = doc.add_paragraph()
                    _set_paragraph_spacing(url_para, 0, 2, line_height)
                    _add_run(url_para, url, font_name, size_detail, italic=True)

    # --- Awards ---
    awards = resume_data.get("awards", [])
    if awards:
        _add_section_header(doc, "Awards", font_name, size_header, section_spacing, divider_style)
        for award in awards:
            if not isinstance(award, dict):
                continue
            award_name = award.get("name", "")
            issuer = award.get("issuer", "")
            date = award.get("date", "")
            description = award.get("description", "")
            if award_name:
                award_para = doc.add_paragraph()
                _set_paragraph_spacing(award_para, 4, 0, 1.0)
                _add_run(award_para, award_name, font_name, size_body, bold=True)
                detail_parts = [p for p in [issuer, date] if p]
                if detail_parts:
                    _add_run(award_para, "  —  " + "  |  ".join(detail_parts), font_name, size_detail)
                if description:
                    desc_para = doc.add_paragraph()
                    _set_paragraph_spacing(desc_para, 0, 2, line_height)
                    _add_run(desc_para, description, font_name, size_body)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _date_range(start: str, end: str) -> str:
    if start and end:
        return f"{start} – {end}"
    return start or end or ""

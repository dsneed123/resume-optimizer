from io import BytesIO

import pytest
from docx import Document

from app.services.docx_export import export_docx
from app.models import default_resume, default_typography


@pytest.fixture
def app(tmp_path):
    from app import create_app
    app = create_app()
    app.instance_path = str(tmp_path)
    app.config['TESTING'] = True
    return app


@pytest.fixture
def ctx(app):
    with app.app_context():
        yield


def test_export_docx_returns_bytes(ctx):
    data = default_resume()
    data['header']['name'] = 'Jane Doe'
    data['header']['email'] = 'jane@example.com'
    result = export_docx(data, default_typography())
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_export_docx_is_valid_docx(ctx):
    data = default_resume()
    data['header']['name'] = 'Jane Doe'
    data['header']['email'] = 'jane@example.com'
    result = export_docx(data, default_typography())
    doc = Document(BytesIO(result))
    assert doc is not None


def test_export_docx_contains_name(ctx):
    data = default_resume()
    data['header']['name'] = 'Jane Doe'
    data['header']['email'] = 'jane@example.com'
    result = export_docx(data, default_typography())
    doc = Document(BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert 'Jane Doe' in full_text


def test_export_docx_with_full_resume(ctx):
    data = default_resume()
    data['header']['name'] = 'John Smith'
    data['header']['email'] = 'john@example.com'
    data['header']['phone'] = '555-1234'
    data['header']['location'] = 'New York, NY'
    data['summary'] = 'Experienced software engineer.'
    data['experience'] = [{
        'company': 'Acme Corp',
        'title': 'Senior Engineer',
        'location': 'New York, NY',
        'start_date': 'Jan 2020',
        'end_date': 'Present',
        'bullets': ['Built scalable APIs', 'Led team of 5'],
    }]
    data['education'] = [{
        'school': 'State University',
        'degree': 'B.S.',
        'field': 'Computer Science',
        'graduation_date': 'May 2018',
        'gpa': '3.8',
        'honors': 'Magna Cum Laude',
    }]
    data['skills'] = [{'category': 'Languages', 'items': ['Python', 'Go', 'TypeScript']}]
    data['certifications'] = [{'name': 'AWS Solutions Architect', 'issuer': 'Amazon', 'date': '2023'}]
    data['projects'] = [{'name': 'MyApp', 'description': 'A web app', 'technologies': 'Python, Flask', 'url': ''}]
    data['awards'] = [{'name': 'Best Engineer', 'issuer': 'Acme', 'date': '2022', 'description': 'Top performer.'}]

    result = export_docx(data, default_typography())
    doc = Document(BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert 'John Smith' in full_text
    assert 'Acme Corp' in full_text
    assert 'State University' in full_text
    assert 'Languages' in full_text
    assert 'AWS Solutions Architect' in full_text
    assert 'MyApp' in full_text
    assert 'Best Engineer' in full_text


def test_export_docx_respects_typography(ctx):
    data = default_resume()
    data['header']['name'] = 'Test User'
    data['header']['email'] = 'test@example.com'
    typo = default_typography()
    typo['font_family'] = 'Georgia'
    typo['font_size_body'] = 11
    result = export_docx(data, typo)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_export_docx_page_size(ctx):
    data = default_resume()
    data['header']['name'] = 'Test User'
    data['header']['email'] = 'test@example.com'
    result = export_docx(data, default_typography())
    doc = Document(BytesIO(result))
    section = doc.sections[0]
    # US Letter: 8.5 x 11 inches (914400 EMU per inch)
    assert abs(section.page_width.inches - 8.5) < 0.01
    assert abs(section.page_height.inches - 11.0) < 0.01


def test_export_docx_empty_sections_omitted(ctx):
    data = default_resume()
    data['header']['name'] = 'Minimal User'
    data['header']['email'] = 'minimal@example.com'
    data['experience'] = []
    data['education'] = []
    data['skills'] = []
    data['certifications'] = []
    data['projects'] = []
    data['awards'] = []
    result = export_docx(data, default_typography())
    doc = Document(BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert 'EXPERIENCE' not in full_text
    assert 'EDUCATION' not in full_text

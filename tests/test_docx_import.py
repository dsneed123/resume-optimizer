import io
import pytest

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from app.services.docx_import import (
    import_docx,
    _classify_section,
    _extract_header_fields,
    _parse_skills_block,
    _parse_experience_block,
    _parse_education_block,
    _parse_certifications_block,
    _parse_projects_block,
    _parse_awards_block,
    _parse_resume_paragraphs,
    _fallback_resume,
    _is_bullet,
    _HEADER_MARKER,
    _extract_degree_and_field,
)

docx_required = pytest.mark.skipif(
    not _DOCX_AVAILABLE,
    reason="python-docx not installed",
)


def _make_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    """Build a minimal DOCX in memory. Each item is (text, style_name)."""
    doc = Document()
    for text, style in paragraphs:
        try:
            p = doc.add_paragraph(text, style=style)
        except Exception:
            p = doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- Unit tests for parsing helpers (no DOCX file needed) ---

def test_classify_section_experience():
    assert _classify_section("EXPERIENCE") == "experience"
    assert _classify_section("Work Experience") == "experience"


def test_classify_section_education():
    assert _classify_section("Education") == "education"


def test_classify_section_skills():
    assert _classify_section("Skills") == "skills"
    assert _classify_section("Technical Skills") == "skills"


def test_classify_section_none():
    assert _classify_section("John Smith") is None
    assert _classify_section("") is None


def test_is_bullet():
    assert _is_bullet("• Built APIs")
    assert _is_bullet("- Led a team")
    assert not _is_bullet("Plain paragraph")


def test_extract_header_fields_basic():
    lines = ["John Smith", "john@example.com", "555-123-4567", "New York, NY"]
    header = _extract_header_fields(lines)
    assert header["name"] == "John Smith"
    assert header["email"] == "john@example.com"
    assert "555" in header["phone"]


def test_extract_header_fields_linkedin():
    lines = ["Jane Doe", "jane@example.com", "linkedin.com/in/janedoe"]
    header = _extract_header_fields(lines)
    assert "linkedin" in header["linkedin"].lower()


def test_parse_skills_block_colon_format():
    lines = ["Languages: Python, Go, TypeScript", "Databases: PostgreSQL, Redis"]
    skills = _parse_skills_block(lines)
    assert len(skills) == 2
    cats = {s["category"] for s in skills}
    assert "Languages" in cats
    assert "Databases" in cats
    lang = next(s for s in skills if s["category"] == "Languages")
    assert "Python" in lang["items"]


def test_parse_skills_block_plain():
    lines = ["Python, JavaScript, React"]
    skills = _parse_skills_block(lines)
    assert len(skills) == 1
    assert "Python" in skills[0]["items"]


def test_parse_skills_block_plain_uses_general_category():
    lines = ["Python, JavaScript, React"]
    skills = _parse_skills_block(lines)
    assert skills[0]["category"] == "General"


def test_parse_skills_block_pipe_separated():
    lines = ["Languages: Python | Go | TypeScript"]
    skills = _parse_skills_block(lines)
    assert len(skills) == 1
    assert skills[0]["category"] == "Languages"
    assert "Go" in skills[0]["items"]
    assert "TypeScript" in skills[0]["items"]


def test_parse_skills_block_bullet_points():
    lines = ["• Python", "• JavaScript", "• React"]
    skills = _parse_skills_block(lines)
    assert len(skills) == 1
    assert skills[0]["category"] == "General"
    items = skills[0]["items"]
    assert "Python" in items
    assert "JavaScript" in items
    assert "React" in items


def test_parse_skills_block_multiple_uncategorized_merged():
    lines = ["Python, Go", "React, Vue"]
    skills = _parse_skills_block(lines)
    assert len(skills) == 1
    assert skills[0]["category"] == "General"
    assert "Python" in skills[0]["items"]
    assert "React" in skills[0]["items"]


def test_parse_skills_block_tools_and_frameworks():
    lines = ["Tools: Git, Docker", "Frameworks: Django, Flask"]
    skills = _parse_skills_block(lines)
    cats = {s["category"] for s in skills}
    assert "Tools" in cats
    assert "Frameworks" in cats
    tools = next(s for s in skills if s["category"] == "Tools")
    assert "Docker" in tools["items"]


def test_parse_experience_block_with_dates():
    lines = [
        "Senior Engineer  |  Acme Corp  |  New York, NY  Jan 2020 - Present",
        "• Built scalable APIs",
        "• Led a team of five engineers",
    ]
    entries = _parse_experience_block(lines)
    assert len(entries) == 1
    assert entries[0]["start_date"].lower().startswith("jan")
    assert entries[0]["end_date"].lower() == "present"
    assert len(entries[0]["bullets"]) == 2


def test_parse_experience_block_year_only_range():
    lines = [
        "Software Engineer  |  Startup Inc  2020-2023",
        "• Shipped product features",
    ]
    entries = _parse_experience_block(lines)
    assert len(entries) == 1
    assert entries[0]["start_date"] == "2020"
    assert entries[0]["end_date"] == "2023"


def test_parse_experience_block_to_separator():
    lines = [
        "Engineer  |  Corp  January 2020 to March 2023",
        "• Did work",
    ]
    entries = _parse_experience_block(lines)
    assert len(entries) == 1
    assert "january" in entries[0]["start_date"].lower()
    assert "march" in entries[0]["end_date"].lower()


def test_parse_experience_block_date_on_separate_line():
    lines = [
        "Software Engineer",
        "Acme Corp",
        "Jan 2020 - Present",
        "• Built APIs",
    ]
    entries = _parse_experience_block(lines)
    assert len(entries) == 1
    assert entries[0]["title"] == "Software Engineer"
    assert entries[0]["company"] == "Acme Corp"
    assert entries[0]["start_date"].lower().startswith("jan")
    assert entries[0]["end_date"].lower() == "present"
    assert len(entries[0]["bullets"]) == 1


def test_parse_experience_block_multiple_positions_bold_header():
    lines = [
        f"{_HEADER_MARKER}Google",
        "Senior Engineer  Jan 2021 - Present",
        "• Led infrastructure work",
        "Junior Engineer  Jan 2019 - Dec 2020",
        "• Built backend services",
    ]
    entries = _parse_experience_block(lines)
    assert len(entries) == 2
    assert entries[0]["company"] == "Google"
    assert entries[0]["title"] == "Senior Engineer"
    assert entries[1]["company"] == "Google"
    assert entries[1]["title"] == "Junior Engineer"


def test_parse_resume_paragraphs_bold_company_header():
    paragraphs = [
        {"text": "EXPERIENCE", "style": "Heading 1", "is_heading": True, "is_list": False, "bold": False, "italic": False},
        {"text": "Google", "style": "Normal", "is_heading": False, "is_list": False, "bold": True, "italic": False},
        {"text": "Senior Engineer  Jan 2021 - Present", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "Built scalable systems", "style": "List Paragraph", "is_heading": False, "is_list": True, "bold": False, "italic": False},
        {"text": "Junior Engineer  Jan 2019 - Dec 2020", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "Wrote backend services", "style": "List Paragraph", "is_heading": False, "is_list": True, "bold": False, "italic": False},
    ]
    result = _parse_resume_paragraphs(paragraphs)
    assert len(result["experience"]) == 2
    assert result["experience"][0]["company"] == "Google"
    assert result["experience"][1]["company"] == "Google"


def test_parse_education_block():
    lines = ["State University, B.S., Computer Science  May 2018", "GPA: 3.8"]
    entries = _parse_education_block(lines)
    assert len(entries) >= 1
    assert entries[0]["school"] == "State University"
    assert entries[0]["gpa"] == "3.8"


def test_parse_education_block_degree_and_field_combined():
    lines = ["MIT  B.S. in Computer Science  May 2021"]
    entries = _parse_education_block(lines)
    assert entries[0]["school"] == "MIT"
    assert entries[0]["degree"] == "B.S."
    assert entries[0]["field"] == "Computer Science"
    assert entries[0]["graduation_date"] == "May 2021"


def test_parse_education_block_expected_date():
    lines = ["State University  B.A. in Economics  Expected May 2025"]
    entries = _parse_education_block(lines)
    assert entries[0]["school"] == "State University"
    assert entries[0]["degree"] == "B.A."
    assert entries[0]["field"] == "Economics"
    assert entries[0]["graduation_date"] == "Expected May 2025"


def test_parse_education_block_full_degree_name():
    lines = ["Stanford University, Bachelor of Science in Computer Science, May 2020"]
    entries = _parse_education_block(lines)
    assert entries[0]["school"] == "Stanford University"
    assert "Bachelor" in entries[0]["degree"]
    assert entries[0]["field"] == "Computer Science"


def test_parse_education_block_phd():
    lines = ["MIT  PhD  Computer Science  2018-2023"]
    entries = _parse_education_block(lines)
    assert entries[0]["school"] == "MIT"
    assert entries[0]["degree"] == "PhD"
    assert entries[0]["field"] == "Computer Science"


def test_extract_degree_and_field_abbreviations():
    assert _extract_degree_and_field("B.S. in Computer Science") == ("B.S.", "Computer Science")
    assert _extract_degree_and_field("MBA") == ("MBA", "")
    assert _extract_degree_and_field("Ph.D. in Physics") == ("Ph.D.", "Physics")
    assert _extract_degree_and_field("M.S. in Data Science") == ("M.S.", "Data Science")


def test_extract_degree_and_field_no_match():
    degree, field = _extract_degree_and_field("Computer Science")
    assert degree == "Computer Science"
    assert field == ""


def test_parse_certifications_block():
    lines = ["AWS Certified Solutions Architect, Amazon, Jan 2023"]
    certs = _parse_certifications_block(lines)
    assert len(certs) == 1
    assert "AWS" in certs[0]["name"]
    assert certs[0]["issuer"] == "Amazon"


def test_parse_projects_block():
    lines = ["Resume Optimizer", "• Built with Flask and React", "Technologies: Python, Flask"]
    projects = _parse_projects_block(lines)
    assert len(projects) == 1
    assert projects[0]["name"] == "Resume Optimizer"
    assert "Flask" in projects[0]["technologies"]


def test_parse_awards_block():
    lines = ["Dean's List, State University, May 2018"]
    awards = _parse_awards_block(lines)
    assert len(awards) == 1
    assert "Dean" in awards[0]["name"]


def test_parse_resume_paragraphs_full():
    paragraphs = [
        {"text": "John Smith", "style": "Normal", "is_heading": False, "is_list": False, "bold": True, "italic": False},
        {"text": "john@example.com", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "555-123-4567", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "SUMMARY", "style": "Heading 1", "is_heading": True, "is_list": False, "bold": False, "italic": False},
        {"text": "Experienced software engineer.", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "EXPERIENCE", "style": "Heading 1", "is_heading": True, "is_list": False, "bold": False, "italic": False},
        {"text": "Senior Engineer  Acme Corp  New York  Jan 2020 - Present", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "Built scalable APIs", "style": "List Paragraph", "is_heading": False, "is_list": True, "bold": False, "italic": False},
        {"text": "EDUCATION", "style": "Heading 1", "is_heading": True, "is_list": False, "bold": False, "italic": False},
        {"text": "State University  B.S.  Computer Science  May 2018", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "SKILLS", "style": "Heading 1", "is_heading": True, "is_list": False, "bold": False, "italic": False},
        {"text": "Languages: Python, Go", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
    ]
    result = _parse_resume_paragraphs(paragraphs)
    assert result["header"]["name"] == "John Smith"
    assert result["header"]["email"] == "john@example.com"
    assert result["summary"] is not None
    assert len(result["experience"]) >= 1
    assert len(result["education"]) >= 1
    assert len(result["skills"]) >= 1


def test_parse_resume_paragraphs_list_items_get_bullet_prefix():
    paragraphs = [
        {"text": "EXPERIENCE", "style": "Heading 1", "is_heading": True, "is_list": False, "bold": False, "italic": False},
        {"text": "Engineer  Acme  Jan 2021 - Dec 2022", "style": "Normal", "is_heading": False, "is_list": False, "bold": False, "italic": False},
        {"text": "Did great work", "style": "List Paragraph", "is_heading": False, "is_list": True, "bold": False, "italic": False},
    ]
    result = _parse_resume_paragraphs(paragraphs)
    assert len(result["experience"]) == 1
    assert len(result["experience"][0]["bullets"]) == 1
    assert result["experience"][0]["bullets"][0] == "Did great work"


def test_fallback_resume_puts_text_in_summary():
    result = _fallback_resume("some raw text")
    assert result["summary"] == "some raw text"
    assert result["experience"] == []
    assert result["education"] == []


def test_fallback_resume_empty_text():
    result = _fallback_resume("")
    assert result["summary"] is None


@docx_required
def test_import_docx_bad_bytes_returns_fallback():
    result = import_docx(b"not a docx file")
    assert isinstance(result, dict)
    assert "header" in result
    assert "experience" in result


@docx_required
def test_import_docx_empty_bytes_returns_fallback():
    result = import_docx(b"")
    assert isinstance(result, dict)
    assert result["experience"] == []


@docx_required
def test_import_docx_real_document():
    content = [
        ("John Smith", "Normal"),
        ("john@example.com", "Normal"),
        ("EXPERIENCE", "Heading 1"),
        ("Senior Engineer  Acme Corp  Jan 2020 - Present", "Normal"),
        ("SKILLS", "Heading 1"),
        ("Languages: Python, Go", "Normal"),
    ]
    file_bytes = _make_docx_bytes(content)
    result = import_docx(file_bytes)
    assert isinstance(result, dict)
    assert "header" in result
    assert "experience" in result
    assert "skills" in result
